import pdfplumber
from docx import Document as DocxDocument
from PIL import Image
import pytesseract
import fitz  # PyMuPDF
import os
from typing import List, Tuple


def extract_text_from_pdf(file_path: str) -> str:
    with pdfplumber.open(file_path) as pdf:
        return "\n".join([page.extract_text() or "" for page in pdf.pages])


def extract_text_from_docx(file_path: str) -> str:
    doc = DocxDocument(file_path)
    return "\n".join([para.text for para in doc.paragraphs])


def extract_text_from_image(file_path: str) -> str:
    image = Image.open(file_path)
    return pytesseract.image_to_string(image)


def extract_text(file_path: str, file_extension: str) -> str:
    if file_extension == "pdf":
        return extract_text_from_pdf(file_path)
    elif file_extension == "docx":
        return extract_text_from_docx(file_path)
    elif file_extension in ["png", "jpg", "jpeg"]:
        return extract_text_from_image(file_path)
    return ""


# ✅ NEW FUNCTION: extract images from PDF and caption with OCR
def extract_images_from_pdf(file_path: str, safe_name: str, output_dir: str) -> List[Tuple[str, str]]:
    doc = fitz.open(file_path)
    image_entries = []

    for i, page in enumerate(doc):
        for img_index, img in enumerate(page.get_images(full=True)):
            xref = img[0]
            base_image = doc.extract_image(xref)
            image_bytes = base_image["image"]
            image_ext = base_image["ext"]
            image_filename = f"{safe_name}_page{i+1}_img{img_index+1}.{image_ext}"
            image_path = os.path.join(output_dir, image_filename)

            with open(image_path, "wb") as f:
                f.write(image_bytes)

            # 🔍 Use OCR to auto-caption
            try:
                image = Image.open(image_path)
                ocr_caption = pytesseract.image_to_string(image).strip()
                context = ocr_caption if ocr_caption else f"image on page {i+1}"
            except Exception:
                context = f"image on page {i+1}"

            image_entries.append((image_filename, context))

    return image_entries


# ✅ NEW FUNCTION: extract images from DOCX and caption with OCR
def extract_images_from_docx(file_path: str, safe_name: str, output_dir: str) -> List[Tuple[str, str]]:
    doc = DocxDocument(file_path)
    image_entries = []

    for i, rel in enumerate(doc.part._rels):
        rel = doc.part._rels[rel]
        if "image" in rel.target_ref:
            image_data = rel.target_part.blob
            image_ext = rel.target_ref.split(".")[-1]
            image_filename = f"{safe_name}_img{i+1}.{image_ext}"
            image_path = os.path.join(output_dir, image_filename)

            with open(image_path, "wb") as f:
                f.write(image_data)

            # 🔍 Use OCR to auto-caption
            try:
                image = Image.open(image_path)
                ocr_caption = pytesseract.image_to_string(image).strip()
                context = ocr_caption if ocr_caption else f"embedded image {i+1}"
            except Exception:
                context = f"embedded image {i+1}"

            image_entries.append((image_filename, context))

    return image_entries
